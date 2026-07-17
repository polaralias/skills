from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

SKILL_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_ROOT = SKILL_ROOT / "scripts"
if str(SCRIPTS_ROOT) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_ROOT))

from build_simple_from_spec import build_document as build_simple_document
from comments_add import add_comments
from comments_extract import extract_comments
from comments_repair import repair_comments
from comments_reply import add_replies
from docx_ooxml import (
    CONTENT_TYPES_PART,
    NS,
    DocxArchive,
    comment_thread_para_id,
    qn,
    rebuild_comment_metadata_parts,
    resolve_relationship_target,
)
from embed_font import embed
from generate_docx import validate_spec
from validate_docx import validate_docx


def make_coauthoring_docx(path: Path, *, include_stray_document: bool = False) -> Path:
    """Rewrite a conventional DOCX into the package shape emitted by co-authoring."""
    source = path.with_name(f"{path.stem}-source.docx")
    document = Document()
    document.add_paragraph("A paragraph that can carry a reviewer comment.")
    document.sections[0].header.paragraphs[0].text = "Co-authoring header"
    document.save(source)

    with zipfile.ZipFile(source, "r") as archive:
        entries = {name: archive.read(name) for name in archive.namelist()}

    document_xml = entries.pop("word/document.xml")
    entries["word/document2.xml"] = document_xml
    document_rels = entries.pop("word/_rels/document.xml.rels")
    rels_root = etree.fromstring(document_rels)
    for rel in rels_root.findall(f"./{{{NS['rels']}}}Relationship"):
        target = rel.get("Target") or ""
        if target and not target.startswith(("/", "http://", "https://")):
            rel.set("Target", f"/{resolve_relationship_target('word/document.xml', target)}")
    entries["word/_rels/document2.xml.rels"] = etree.tostring(
        rels_root,
        encoding="utf-8",
        xml_declaration=True,
    )

    package_rels = etree.fromstring(entries["_rels/.rels"])
    office_document_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"
    for rel in package_rels.findall(f"./{{{NS['rels']}}}Relationship"):
        if rel.get("Type") == office_document_type:
            rel.set("Target", "word/document2.xml")
    entries["_rels/.rels"] = etree.tostring(package_rels, encoding="utf-8", xml_declaration=True)

    content_types = etree.fromstring(entries[CONTENT_TYPES_PART])
    for override in content_types.findall(f"./{{{NS['ct']}}}Override"):
        if override.get("PartName") == "/word/document.xml":
            override.set("PartName", "/word/document2.xml")
    etree.SubElement(
        content_types,
        f"{{{NS['ct']}}}Default",
        {"Extension": "dat", "ContentType": "application/octet-stream"},
    )
    entries[CONTENT_TYPES_PART] = etree.tostring(content_types, encoding="utf-8", xml_declaration=True)
    entries["[trash]/0000000000000000.dat"] = b"orphaned co-authoring residue"
    if include_stray_document:
        entries["word/document.xml"] = b"<not-the-authoritative-document/>"

    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, data in entries.items():
            archive.writestr(name, data)
    source.unlink()
    return path


def test_spec_rejects_raw_ooxml_without_flag() -> None:
    spec = {
        "title": "Raw spec",
        "blocks": [{"type": "raw_docx_xml", "xml": "<w:p/>", "reason": "test"}],
    }

    try:
        validate_spec(spec, allow_raw_ooxml=False)
    except SystemExit as exc:
        assert "raw_docx_xml blocks require --allow_raw_ooxml" in str(exc)
    else:
        raise AssertionError("Expected SystemExit for raw_docx_xml without allow flag")


def test_simple_builder_creates_docx(tmp_path: Path) -> None:
    output_path = tmp_path / "simple.docx"
    spec = {
        "title": "Simple output",
        "subtitle": "Subtitle",
        "blocks": [
            {"type": "paragraph", "text": "Hello world."},
            {"type": "bullets", "items": ["One", "Two"]},
        ],
    }

    build_simple_document(spec, output_path)

    assert output_path.exists()
    document = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Simple output" in text
    assert "Hello world." in text


def test_embed_is_noop_without_configured_font(tmp_path: Path, monkeypatch) -> None:
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("No custom font configured.")
    document.save(input_path)

    monkeypatch.delenv("DOCX_THEME_PRIMARY_FONT_TTF", raising=False)
    result = embed(str(input_path), str(output_path))

    assert Path(result) == output_path
    assert output_path.exists()
    assert output_path.read_bytes() == input_path.read_bytes()


def test_main_part_and_story_parts_resolve_document2(tmp_path: Path) -> None:
    input_path = make_coauthoring_docx(tmp_path / "coauthoring.docx", include_stray_document=True)
    archive = DocxArchive.load(input_path)

    assert archive.main_document_part() == "word/document2.xml"
    assert archive.main_document_rels_part() == "word/_rels/document2.xml.rels"
    assert archive.story_parts()[0] == "word/document2.xml"
    assert "word/document.xml" not in archive.story_parts()
    assert "word/header1.xml" in archive.story_parts()
    assert archive.document_root().tag == qn("w", "document")


def test_normalise_main_document_part_renames_and_cleans(tmp_path: Path) -> None:
    input_path = make_coauthoring_docx(tmp_path / "coauthoring.docx", include_stray_document=True)
    archive = DocxArchive.load(input_path)

    assert archive.normalise_main_document_part() is True
    assert archive.main_document_part() == "word/document.xml"
    assert archive.has("word/document.xml")
    assert archive.document_root().tag == qn("w", "document")
    assert not archive.has("word/document2.xml")
    assert archive.has("word/_rels/document.xml.rels")
    assert not archive.has("word/_rels/document2.xml.rels")
    assert not any(name.startswith("[trash]/") for name in archive.entries)

    package_rels = archive.read_xml("_rels/.rels")
    office_document_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"
    office_rel = next(
        rel
        for rel in package_rels.findall(f"./{{{NS['rels']}}}Relationship")
        if rel.get("Type") == office_document_type
    )
    assert office_rel.get("Target") == "word/document.xml"
    assert all(
        not (rel.get("Target") or "").startswith("/word/")
        for rel in archive.read_xml("word/_rels/document.xml.rels").findall(f"./{{{NS['rels']}}}Relationship")
        if rel.get("TargetMode") != "External"
    )
    assert archive.normalise_main_document_part() is False


def test_comment_reply_roundtrip_on_coauthoring_doc(tmp_path: Path) -> None:
    input_path = make_coauthoring_docx(tmp_path / "coauthoring.docx")
    commented_path = tmp_path / "commented.docx"
    replied_path = tmp_path / "replied.docx"

    add_comments(
        input_path,
        commented_path,
        [("paragraph that can carry", "Please update this paragraph.")],
        author="Reviewer",
        initials="RV",
        ignore_case=False,
    )
    add_replies(
        commented_path,
        {"0": "Updated."},
        replied_path,
        author="Author",
        initials="AU",
    )

    archive = DocxArchive.load(replied_path)
    report = extract_comments(replied_path)
    validation = validate_docx(replied_path)
    assert archive.main_document_part() == "word/document.xml"
    assert report["comment_count"] == 2
    assert validation["ok"], validation["errors"]


def test_validate_passes_on_coauthoring_doc(tmp_path: Path) -> None:
    input_path = make_coauthoring_docx(tmp_path / "coauthoring.docx")

    validation = validate_docx(input_path)

    assert validation["ok"], validation["errors"]


def test_validate_warns_about_external_relationships(tmp_path: Path) -> None:
    input_path = tmp_path / "external-link.docx"
    document = Document()
    document.add_paragraph("External content test")
    document.save(input_path)

    archive = DocxArchive.load(input_path)
    rels = archive.relationships_root(archive.main_document_part())
    etree.SubElement(
        rels,
        qn("rels", "Relationship"),
        {
            "Id": "rIdExternal",
            "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            "Target": "https://example.invalid/tracker",
            "TargetMode": "External",
        },
    )
    archive.save_relationships_root(archive.main_document_part(), rels)
    archive.write(input_path)

    validation = validate_docx(input_path)

    assert validation["ok"], validation["errors"]
    assert any("targets external content" in warning for warning in validation["warnings"])


def test_comment_repair_normalises_coauthoring_package_without_comments(tmp_path: Path) -> None:
    input_path = make_coauthoring_docx(tmp_path / "coauthoring.docx")
    repaired_path = tmp_path / "repaired.docx"

    stats = repair_comments(input_path, repaired_path)

    archive = DocxArchive.load(repaired_path)
    assert stats["comments_present"] == 0
    assert archive.main_document_part() == "word/document.xml"
    assert not any(name.startswith("[trash]/") for name in archive.entries)
    validation = validate_docx(repaired_path)
    assert validation["ok"], validation["errors"]


def test_thread_para_id_prefers_existing_commentex_key() -> None:
    comment = etree.Element(qn("w", "comment"), {qn("w", "id"): "7"})
    etree.SubElement(comment, qn("w", "p"), {qn("w14", "paraId"): "11111111"})
    etree.SubElement(comment, qn("w", "p"), {qn("w14", "paraId"): "22222222"})
    comments_extended = etree.Element(qn("w15", "commentsEx"))
    etree.SubElement(
        comments_extended,
        qn("w15", "commentEx"),
        {qn("w15", "paraId"): "22222222", qn("w15", "done"): "0"},
    )

    assert comment_thread_para_id(comment, comments_extended) == "22222222"


def test_thread_para_id_falls_back_to_first_paragraph() -> None:
    comment = etree.Element(qn("w", "comment"), {qn("w", "id"): "7"})
    etree.SubElement(comment, qn("w", "p"), {qn("w14", "paraId"): "11111111"})
    etree.SubElement(comment, qn("w", "p"), {qn("w14", "paraId"): "22222222"})

    assert comment_thread_para_id(comment, etree.Element(qn("w15", "commentsEx"))) == "11111111"


def test_metadata_rebuild_preserves_later_key_and_only_ensures_first_reference(tmp_path: Path) -> None:
    source_path = tmp_path / "source.docx"
    commented_path = tmp_path / "commented.docx"
    document = Document()
    document.add_paragraph("A paragraph that can carry a reviewer comment.")
    document.save(source_path)
    add_comments(
        source_path,
        commented_path,
        [("paragraph that can carry", "First paragraph of comment.")],
        author="Reviewer",
        initials="RV",
        ignore_case=False,
    )

    archive = DocxArchive.load(commented_path)
    comments_root = archive.comments_root()
    comment = comments_root.find(f"./{qn('w', 'comment')}")
    assert comment is not None
    second_para_id = "22222222"
    second_paragraph = etree.SubElement(
        comment,
        qn("w", "p"),
        {
            qn("w14", "paraId"): second_para_id,
            qn("w14", "textId"): "33333333",
            qn("w", "rsidR"): "44444444",
            qn("w", "rsidRDefault"): "44444444",
        },
    )
    run = etree.SubElement(second_paragraph, qn("w", "r"), {qn("w", "rsidR"): "44444444"})
    etree.SubElement(run, qn("w", "t")).text = "Continuation paragraph."

    comments_extended = archive.comments_extended_root()
    comment_ex = comments_extended.find(f"./{qn('w15', 'commentEx')}")
    assert comment_ex is not None
    original_para_id = comment_ex.get(qn("w15", "paraId"))
    comment_ex.set(qn("w15", "paraId"), second_para_id)
    comments_ids = archive.comments_ids_root()
    comment_id = comments_ids.find(f"./{qn('w16cid', 'commentId')}")
    assert comment_id is not None
    assert comment_id.get(qn("w16cid", "paraId")) == original_para_id
    comment_id.set(qn("w16cid", "paraId"), second_para_id)
    archive.save_comments_root(comments_root)
    archive.save_comments_extended_root(comments_extended)
    archive.save_comments_ids_root(comments_ids)

    rebuild_comment_metadata_parts(archive, comments_root)

    rebuilt_comment_ex = archive.comments_extended_root().find(f"./{qn('w15', 'commentEx')}")
    rebuilt_comment_id = archive.comments_ids_root().find(f"./{qn('w16cid', 'commentId')}")
    assert rebuilt_comment_ex is not None
    assert rebuilt_comment_id is not None
    assert rebuilt_comment_ex.get(qn("w15", "paraId")) == second_para_id
    assert rebuilt_comment_id.get(qn("w16cid", "paraId")) == second_para_id
    paragraphs = comment.findall(f"./{qn('w', 'p')}")
    assert paragraphs[0].find(f".//{qn('w', 'annotationRef')}") is not None
    assert paragraphs[1].find(f".//{qn('w', 'annotationRef')}") is None
    rebuilt_path = tmp_path / "rebuilt.docx"
    archive.write(rebuilt_path)
    validation = validate_docx(rebuilt_path)
    assert validation["ok"], validation["errors"]
    assert not any("annotation reference run" in warning for warning in validation["warnings"])
