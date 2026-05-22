#!/usr/bin/env python3
"""Render a DOCX JSON spec through the fast simple path."""

from __future__ import annotations

import argparse
import json
import os
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
_LOGO_PATH_VALUE = (
    os.environ.get("DOCX_THEME_LOGO_PROTECTED_PATH")
    or os.environ.get("DOCX_THEME_LOGO_PATH")
)
LOGO_PATH = Path(_LOGO_PATH_VALUE).expanduser().resolve() if _LOGO_PATH_VALUE else None
ACCENT_ORANGE = RGBColor(0xFF, 0x45, 0x00)
ACCENT_BLUE = RGBColor(0x2F, 0x5B, 0xFF)
TEXT_BLACK = RGBColor(0x22, 0x22, 0x22)
TEXT_GREY = RGBColor(0x55, 0x55, 0x55)


def set_bottom_border(paragraph, color: str = "FF4500", size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT_BLACK

    for style_name, size, color in (
        ("Title", 22, TEXT_BLACK),
        ("Heading 1", 16, ACCENT_BLUE),
        ("Heading 2", 13, TEXT_BLACK),
        ("Heading 3", 11.5, TEXT_BLACK),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_header(document: Document) -> None:
    if LOGO_PATH is None or not LOGO_PATH.is_file():
        return
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.alignment = 0
    paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(1.25))


def add_title(document: Document, spec: dict[str, Any]) -> None:
    title = document.add_paragraph(style="Title")
    title.add_run(spec["title"])
    set_bottom_border(title)
    subtitle = spec.get("subtitle") or spec.get("customer")
    if subtitle:
        para = document.add_paragraph()
        run = para.add_run(subtitle)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = TEXT_GREY


def add_paragraph(document: Document, block: dict[str, Any]) -> None:
    para = document.add_paragraph()
    run = para.add_run(block["text"])
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    if block.get("emphasis") == "muted":
        run.font.color.rgb = TEXT_GREY
    elif block.get("emphasis") == "strong":
        run.bold = True


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        para = document.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)


def add_callout(document: Document, block: dict[str, Any]) -> None:
    label = block["variant"].replace("_", " ").upper()
    title_text = block.get("title")
    para = document.add_paragraph()
    label_run = para.add_run(label)
    label_run.bold = True
    label_run.font.name = "Arial"
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = ACCENT_BLUE if block["variant"] != "warning" else ACCENT_ORANGE
    if title_text:
        title_run = para.add_run(f": {title_text}")
        title_run.bold = True
        title_run.font.name = "Arial"
        title_run.font.size = Pt(10)
    for child in block["body"]:
        if child["type"] == "paragraph":
            add_paragraph(document, child)
        elif child["type"] == "bullets":
            add_bullets(document, child["items"])
        else:
            raise ValueError(f"Unsupported callout child block: {child['type']}")


def add_table(document: Document, block: dict[str, Any]) -> None:
    headers = block["headers"]
    rows = block["rows"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        hdr_cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row[: len(headers)]):
            cells[index].text = value


def render_block(document: Document, block: dict[str, Any]) -> None:
    block_type = block["type"]
    if block_type == "paragraph":
        add_paragraph(document, block)
    elif block_type == "heading":
        document.add_heading(block["text"], level=block.get("level", 1))
    elif block_type == "bullets":
        add_bullets(document, block["items"])
    elif block_type == "section_banner":
        if block.get("label"):
            label = document.add_paragraph()
            run = label.add_run(block["label"])
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            run.font.color.rgb = ACCENT_BLUE
        document.add_heading(block["title"], level=1)
        if block.get("subtitle"):
            add_paragraph(document, {"type": "paragraph", "text": block["subtitle"], "emphasis": "muted"})
    elif block_type == "callout":
        add_callout(document, block)
    elif block_type == "table":
        add_table(document, block)
    elif block_type == "page_break":
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    elif block_type == "raw_docx_xml":
        raise ValueError("raw_docx_xml is not supported by the simple renderer")
    else:
        raise ValueError(f"Unknown block type: {block_type}")


def build_document(spec: dict[str, Any], output_path: Path) -> Path:
    document = Document()
    style_document(document)
    add_header(document)
    add_title(document, spec)
    for block in spec["blocks"]:
        render_block(document, block)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a simple DOCX from a JSON spec.")
    parser.add_argument("--input", required=True, help="Input spec JSON path")
    parser.add_argument("--output", required=True, help="Destination path for the DOCX result")
    args = parser.parse_args()

    spec = json.loads(Path(args.input).read_text(encoding="utf-8"))
    print(build_document(spec, Path(args.output).expanduser().resolve()))


if __name__ == "__main__":
    main()
